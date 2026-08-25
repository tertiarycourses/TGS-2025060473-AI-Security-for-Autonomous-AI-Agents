#!/usr/bin/env python3
"""WSQ assessment set builder — AI Security for Autonomous AI Agents (TGS-2025060473).

Produces four DOCX in ../assessment/:
  1. Written Assessment (SAQ) — Question Paper      (5 open-ended SAQ, one per K statement)
  2. Written Assessment (SAQ) — Answer Key
  3. Case Study — Question Paper                    (3 open-ended questions, one per LO)
  4. Case Study — Answer Key

Layout contract (WSQ QA):
  page 1 = cover page naming the instrument exactly
  page 2 = Trainee Information + Instructions + Grading block   (question papers)
  page 3 = scenario / questions start here
Answer keys carry the cover page + K/A coverage mapping table.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
OUT = os.path.join(REPO, "assessment")
ASSETS = os.path.join(HERE, "assets")
ORG_LOGO = os.path.join(ASSETS, "tertiary-infotech-logo.png")
COURSE_LOGO = os.path.join(ASSETS, "n8n-course-logo.png")

# ------------------------------------------------------------------ course facts
TITLE = "AI Security for Autonomous AI Agents"
COURSE_CODE = "TGS-2025060473"
TSC_TITLE = "Generative AI Principles and Applications"
TSC_CODE = "ICT-INT-0052-1.1"
ORG = "Tertiary Infotech Pte Ltd"
UEN = "UEN: 201200696W"
VERSION = "4.3"
VERSION_DATE = "25 August 2026"
LMS_URL = "https://lms-tms.tertiaryinfotech.com/"
COPYRIGHT = ("This material belongs to Tertiary Infotech Pte Ltd (UEN: 201200696W). "
             "All Rights Reserved.")

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

K_STATEMENTS = {
    "K1": "Importance of data quality, preprocessing, model pipeline and model training "
          "(e.g., impact of data bias from training data)",
    "K2": "Underlying principles, core concepts and theories governing generative AI",
    "K3": "Difference between generative and discriminative models",
    "K4": "Impact of prompt engineering on the model outputs of generative AI",
    "K5": "Generative AI model workings, including training data, algorithms, and outputs",
}
A_STATEMENTS = {
    "A1": "Analyse limitations and potential biases in AI-generated content",
    "A2": "Identify the ethical implications and societal impact of AI-generated content",
    "A3": "Apply understanding of generative AI principles to use cases",
    "A4": "Demonstrate the use of generation AI in diverse applications (e.g., summarisation, "
          "inference, reasoning, transformation of content, augmentation of content)",
    "A5": "Analyse generative AI models' performance metrics and evaluate the influence of "
          "prompt variations",
}
LO_STATEMENTS = {
    "LO1": "Demonstrate generative AI concepts and applications relevant to customer service "
           "and hospitality management.",
    "LO2": "Apply prompt engineering techniques and analyse output variations to improve "
           "generative AI performance in service settings.",
    "LO3": "Identify ethical risks and analyse bias in AI-generated content used in customer "
           "engagement.",
}


# ================================================================== docx helpers
def _field(paragraph, instr, default=""):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = default
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (b, it, sep, t, end):
        run._r.append(el)
    return run


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def _fix_widths(table, widths_in):
    """Disable autofit and pin column widths, so Word/LibreOffice honour them."""
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)
    for i, w in enumerate(widths_in):
        if i < len(table.columns):
            table.columns[i].width = Inches(w)


def _cell_text(cell, text, bold=False, size=9.5, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"


def new_doc():
    d = Document()
    n = d.styles["Normal"]
    n.font.name = "Arial"; n.font.size = Pt(11)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for sec in d.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    return d


def line(d, text="", bold=False, size=11, color=DARK, after=6, before=0, align=None,
         italic=False, indent=None):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p


def runs(d, segments, after=6, style=None, size=11, indent=None):
    p = d.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold in segments:
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p


def bullet(d, text, size=10.5, indent=0.25):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(indent + 0.25)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p


def page_break(d):
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def keep_with_next(paragraph):
    """Bind a paragraph to the one that follows so a heading never orphans."""
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    return paragraph


def keep_lines(paragraph):
    """Stop a single paragraph being split across a page boundary."""
    paragraph.paragraph_format.keep_together = True
    return paragraph


def img(d, path, width_in):
    if path and os.path.exists(path):
        p = d.add_paragraph(); p.alignment = AL.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(path, width=Inches(width_in))
        return True
    return False


def footer_block(d):
    sec = d.sections[0]
    f = sec.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.text = ""
    p.alignment = AL.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(COPYRIGHT)
    r.font.size = Pt(7.5); r.font.color.rgb = GREY; r.font.name = "Arial"
    p2 = f.add_paragraph()
    p2.alignment = AL.CENTER
    r2 = p2.add_run("Page ")
    r2.font.size = Pt(8.5); r2.font.color.rgb = GREY; r2.font.name = "Arial"
    _field(p2, "PAGE", "1").font.size = Pt(8.5)
    r3 = p2.add_run(" of ")
    r3.font.size = Pt(8.5); r3.font.color.rgb = GREY; r3.font.name = "Arial"
    _field(p2, "NUMPAGES", "1").font.size = Pt(8.5)


def enable_update_fields(d):
    settings = d.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        settings.remove(existing)
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
    settings.insert_element_before(uf, "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat")


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    nr = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1F6FEB"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), "Arial"); rf.set(qn("w:hAnsi"), "Arial")
    rPr.append(rf)
    nr.append(rPr)
    t = OxmlElement("w:t"); t.text = text; nr.append(t)
    hl.append(nr); paragraph._p.append(hl)


# ================================================================== PAGE 1 — cover
def cover_page(d, instrument, doc_kind):
    """instrument: 'Written Assessment (SAQ)' or 'Case Study'.
    doc_kind: 'Question Paper' or 'Answer Key'."""
    d.add_paragraph()
    img(d, ORG_LOGO, 2.1)
    line(d, ORG, bold=True, size=13, color=DARK, before=2, after=1, align=AL.CENTER)
    line(d, UEN, size=10, color=GREY, after=14, align=AL.CENTER)

    line(d, "WSQ ASSESSMENT", size=11, color=GREY, after=4, align=AL.CENTER)
    line(d, instrument.upper(), bold=True, size=24, color=BRAND, after=2, align=AL.CENTER)
    line(d, doc_kind.upper(), bold=True, size=14, color=DARK, after=12, align=AL.CENTER)

    line(d, "For", size=11, color=GREY, after=6, align=AL.CENTER)
    img(d, COURSE_LOGO, 1.0)
    line(d, TITLE, bold=True, size=19, color=DARK, before=4, after=3, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}", size=12, color=DARK, after=2, align=AL.CENTER)
    line(d, f"TSC: {TSC_TITLE} ({TSC_CODE})", size=11, color=GREY, after=14, align=AL.CENTER)

    line(d, "Conducted by", size=10.5, color=GREY, after=2, align=AL.CENTER)
    line(d, ORG, bold=True, size=13, color=DARK, after=2, align=AL.CENTER)
    line(d, UEN, size=10.5, color=GREY, after=12, align=AL.CENTER)

    line(d, f"Version {VERSION}", bold=True, size=12, color=BRAND, after=2, align=AL.CENTER)
    line(d, VERSION_DATE, size=10.5, color=GREY, after=0, align=AL.CENTER)
    page_break(d)


# ================================================================== PAGE 2 — admin
def admin_page(d, instrument, duration, total_marks, n_items, item_word, instructions):
    """Page 2 of every question paper. Must fit on ONE page: Trainee Information,
    Instructions and the Grading block, so that the questions start on page 3."""
    line(d, instrument, bold=True, size=13, color=BRAND, after=1, align=AL.CENTER)
    line(d, TITLE, bold=True, size=11, color=DARK, after=1, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}  ·  TSC: {TSC_CODE}  ·  Version {VERSION}",
         size=8.5, color=GREY, after=7, align=AL.CENTER)

    # ------------------------------------------------ A: Trainee Information
    line(d, "SECTION A: TRAINEE INFORMATION", bold=True, size=10.5, color=BRAND, after=3)
    t = d.add_table(rows=0, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = [
        ("Trainee Name (as per NRIC/FIN)", "NRIC/FIN (last 4 characters)"),
        ("Date of Assessment", "Trainer Name"),
        ("Assessor Name", ""),
    ]
    for left, right in grid:
        c = t.add_row().cells
        _cell_text(c[0], left, bold=True, size=8.5)
        _shade(c[0], "F2F5FA")
        _cell_text(c[1], "", size=9)
        _cell_text(c[2], right, bold=True, size=8.5)
        if right:
            _shade(c[2], "F2F5FA")
        _cell_text(c[3], "", size=9)
    widths = [1.68, 1.62, 1.68, 1.62]
    _fix_widths(t, widths)
    for r in t.rows:
        r.height = Cm(0.62)
    line(d, "", after=6)

    # ------------------------------------------------ Assessment summary strip
    st = d.add_table(rows=2, cols=4)
    st.style = "Table Grid"
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ["Instrument", "No. of Questions", "Duration", "Total Marks"]
    vals = [instrument, f"{n_items} {item_word}", duration, f"{total_marks} marks"]
    for i, h in enumerate(heads):
        _cell_text(st.rows[0].cells[i], h, bold=True, size=8.5, color=WHITE, align=AL.CENTER)
        _shade(st.rows[0].cells[i], "1F6FEB")
    for i, v in enumerate(vals):
        _cell_text(st.rows[1].cells[i], v, size=8.5, align=AL.CENTER)
    _fix_widths(st, widths)
    line(d, "", after=6)

    # ------------------------------------------------ B: Instructions
    line(d, "SECTION B: INSTRUCTIONS TO CANDIDATE", bold=True, size=10.5, color=BRAND, after=3)
    ib = d.add_table(rows=1, cols=1)
    ib.style = "Table Grid"
    cell = ib.rows[0].cells[0]
    cell.text = ""
    first = True
    for txt in instructions:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt)
        r.font.size = Pt(8.5); r.font.name = "Arial"; r.font.color.rgb = DARK
    # LMS submission line with a real hyperlink
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Submission: complete your answers on this document and upload the completed "
                  "paper to the Learning Management System at ")
    r.font.size = Pt(8.5); r.font.name = "Arial"; r.font.color.rgb = DARK
    add_hyperlink(p, LMS_URL, LMS_URL)
    line(d, "", after=6)

    # ------------------------------------------------ C: Grading
    line(d, "SECTION C: GRADING (FOR OFFICIAL USE ONLY)", bold=True, size=10.5,
         color=BRAND, after=3)
    line(d, "The candidate is graded COMPETENT (C) only when ALL questions are assessed as "
            "satisfactory against the marking guide. Any question assessed as unsatisfactory "
            "results in NOT YET COMPETENT (NYC) and the candidate is offered re-assessment.",
         size=8.5, color=GREY, after=4)
    g = d.add_table(rows=3, cols=3)
    g.style = "Table Grid"
    g.alignment = WD_TABLE_ALIGNMENT.CENTER
    gw = [1.8, 2.4, 2.4]

    _cell_text(g.rows[0].cells[0], "Overall Result", bold=True, size=9)
    _shade(g.rows[0].cells[0], "F2F5FA")
    _cell_text(g.rows[0].cells[1], "☐   COMPETENT (C)", bold=True, size=9, align=AL.CENTER)
    _cell_text(g.rows[0].cells[2], "☐   NOT YET COMPETENT (NYC)", bold=True, size=9,
               align=AL.CENTER)

    _cell_text(g.rows[1].cells[0], "Assessor Feedback", bold=True, size=9)
    _shade(g.rows[1].cells[0], "F2F5FA")
    g.rows[1].cells[1].merge(g.rows[1].cells[2])
    g.rows[1].height = Cm(1.15)

    _cell_text(g.rows[2].cells[0], "Assessor Signature / Date", bold=True, size=9)
    _shade(g.rows[2].cells[0], "F2F5FA")
    g.rows[2].cells[1].merge(g.rows[2].cells[2])
    g.rows[2].height = Cm(0.75)

    _fix_widths(g, gw)
    page_break(d)


# ================================================================== answer box
def answer_box(d, lines_n=8, caption="Answer"):
    """A bordered answer box that can NEVER fragment across a page break.

    The box is a single-row, single-cell table. Word/LibreOffice will split a table
    row across pages unless cantSplit is set — which is exactly the failure mode
    (a box sliced open with the footer printing through it). We set cantSplit so
    the row moves whole, and the caller sizes lines_n so the box fits its page.
    """
    t = d.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    row = t.rows[0]
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)

    cell = row.cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(caption + ":")
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY; r.font.name = "Arial"
    for _ in range(lines_n):
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(10)
        rr = pp.add_run("")
        rr.font.size = Pt(11); rr.font.name = "Arial"
    line(d, "", after=4)


# ================================================================== WA content
# (qno, K code, marks, context, question stem, [sub-prompts], answer_lines)
WRITTEN = [
    (1, "K2", 8,
     "Topic 1 — Generative AI, Agentic AI and AI Agents",
     "During the course you saw how generative AI writes: an autoregressive large language model "
     "reads your context as tokens and predicts the next token one at a time. You also saw the "
     "same model wrapped in an agentic loop of context, plan, execute and verify.",
     ["(a) Explain the underlying principles and core concepts of how a generative AI model "
      "produces its output, covering tokens, next-token prediction and why the same prompt can "
      "give different answers. (4 marks)",
      "(b) Explain what the agentic loop adds on top of the model, and why an agent is therefore "
      "a system rather than just a model. (4 marks)"],
     11),
    (2, "K3", 8,
     "Topic 1 — Generative vs discriminative models",
     "A customer-service team uses AI in two ways: one model DRAFTS replies to guests, and a "
     "second model CLASSIFIES each incoming message as a complaint, a booking or a general "
     "enquiry so it can be routed.",
     ["(a) State the difference between a generative model and a discriminative model, covering "
      "what each one learns and what each one produces. (4 marks)",
      "(b) Identify which model type performs each role above, and give ONE further example of "
      "each type from a customer-service or hospitality setting. (4 marks)"],
     11),
    (3, "K1", 8,
     "Topic 3 — Data quality, training data and bias",
     "An AI agent is given access to a hotel's booking spreadsheet and past guest correspondence, "
     "and it can read, summarise and edit that data. Some of the historical data is incomplete or "
     "skewed toward one type of guest.",
     ["(a) Explain why data quality, preprocessing, the model pipeline and model training matter "
      "to how a generative AI system behaves, including how bias in the data can reach the "
      "outputs. (4 marks)",
      "(b) Identify TWO problems that poor-quality or biased data could cause in THIS deployment, "
      "and state one way to reduce each. (4 marks)"],
     11),
    (4, "K4", 8,
     "Topic 2 — Prompt engineering and output variation",
     "In the Topic 2 activities you sent the SAME data to an AI agent using a weak prompt (for "
     "example, \"analyse this excel\") and then a strong prompt that named a role, a task and a "
     "format.",
     ["(a) Explain the impact of prompt engineering on the outputs of generative AI, covering how "
      "a well-formed prompt improves the output. (4 marks)",
      "(b) Using an example from your activity work, describe how changing the prompt changed the "
      "output, and state TWO principles of a good prompt. (4 marks)"],
     11),
    (5, "K5", 8,
     "Topic 1 & 3 — How a model works and what an agent adds",
     "The same generative model is deployed first as a chatbot that only answers, and then as an "
     "autonomous agent that is given tools to send messages and edit records. Nothing about the "
     "model itself changed.",
     ["(a) Explain how a generative AI model works — its training data, its algorithm at "
      "inference time, and its outputs. (4 marks)",
      "(b) Explain why the same model becomes more risky once its outputs can trigger tool "
      "actions, and name TWO controls that reduce that risk. (4 marks)"],
     11),
]

WA_ANSWERS = {
    1: dict(
        source="Topic 1 — How generative AI works (autoregressive LLM); the agentic loop; "
               "generative vs agentic vs agents summary.",
        points=[
            ("(a) How a generative model produces output — 4 marks", [
                "A large language model reads the context (system prompt, history and your "
                "request) as tokens. (1 mark)",
                "It predicts the NEXT token from the probabilities it learned, then appends it "
                "and repeats — writing one token at a time. (1 mark)",
                "It is autoregressive: each token depends on everything before it. (1 mark)",
                "Because it SAMPLES from likely tokens (and temperature/context vary), the same "
                "prompt can give different answers — it is not a fixed database lookup. (1 mark)",
            ]),
            ("(b) What the agentic loop adds — 4 marks", [
                "The loop runs context, plan, execute and verify, repeating until the goal is "
                "met. (1 mark)",
                "'Execute' lets the model call tools and take real actions, not just answer. "
                "(1 mark)",
                "It keeps state and can retry or re-plan when verification fails. (1 mark)",
                "So an agent is a whole running SYSTEM — model plus loop, tools and identity — "
                "not just a model. (1 mark)",
            ]),
        ],
        guidance="Award (a) for the mechanism in plain words: 'it predicts the next word based on "
                 "everything so far, one at a time' earns marks without the term 'autoregressive'. "
                 "For (b) the mark is for the LOOP and TOOLS making it a system; 'it is smarter' "
                 "alone scores 0. Accept 'generation is a step, agency is a loop, an agent is a "
                 "system'.",
    ),
    2: dict(
        source="Topic 1 — Generative vs discriminative models; strengths and limits for service work.",
        points=[
            ("(a) The difference — 4 marks", [
                "Generative model: learns the patterns in data and PRODUCES new content — text, "
                "images, audio, code. (2 marks: 1 for what it learns, 1 for what it produces.)",
                "Discriminative model: learns to SORT or SCORE existing input into classes — for "
                "example complaint vs enquiry — and produces a label or score, not new content. "
                "(2 marks: 1 for what it learns, 1 for what it produces.)",
            ]),
            ("(b) Roles above + one further example each — 4 marks", [
                "The model that DRAFTS replies is the generative model. (1 mark)",
                "The model that CLASSIFIES incoming messages is the discriminative model. "
                "(1 mark)",
                "A further generative example: writing a marketing caption, summarising a thread, "
                "translating a message. (1 mark)",
                "A further discriminative example: spam/not-spam, sentiment positive/negative, "
                "fraud/not-fraud, routing a ticket to a department. (1 mark)",
            ]),
        ],
        guidance="Any reasonable service example is acceptable. Do not require the words "
                 "'generative'/'discriminative' if the candidate clearly describes 'makes new "
                 "content' vs 'sorts/scores'. A candidate who swaps the two roles in (b) scores 0 "
                 "for that part.",
    ),
    3: dict(
        source="Topic 3 — AI data governance; new principles (provenance, traceability); "
               "Activity 5 (data governance policy).",
        points=[
            ("(a) Why data quality / preprocessing / pipeline / training matter — 4 marks", [
                "A generative model has no ground truth of its own — its outputs reflect the data "
                "it learned from and the data it is given. (1 mark)",
                "Poor preprocessing lets incomplete, duplicated or mislabelled data through, which "
                "shapes the answers. (1 mark)",
                "Bias in the data becomes bias in the OUTPUT, stated with the confidence of a "
                "correct answer. (1 mark)",
                "Because one dataset drives many answers, a data-quality problem is amplified at "
                "scale. (1 mark)",
            ]),
            ("(b) Two problems in THIS deployment + a fix each — 4 marks (2 each)", [
                "Skewed history — the agent gives better answers for the common guest type and "
                "worse ones for under-represented guests. Fix: check accuracy by segment and "
                "broaden or rebalance the data.",
                "Incomplete/incorrect records — the agent summarises or edits from wrong data and "
                "produces wrong output. Fix: validate and clean data, and require human approval "
                "before the agent changes records.",
                "[Also accept: no provenance on old records, so bad data cannot be traced — fix "
                "with a data inventory and logging; personal data used beyond its purpose — fix "
                "with purpose limitation and minimisation.]",
            ]),
        ],
        guidance="(b) needs TWO problems tied to THIS hotel deployment plus a fix each. A generic "
                 "'garbage in, garbage out' with no deployment link caps at 2 of 4 for (b).",
    ),
    4: dict(
        source="Topic 2 — Prompt engineering principles; good vs bad prompts; Activities 2-4.",
        points=[
            ("(a) Impact of prompt engineering on outputs — 4 marks", [
                "The prompt is the main way you control the output. (1 mark)",
                "A well-formed prompt naming a role, task, context, format and constraints raises "
                "accuracy and usefulness. (1 mark)",
                "The model conditions on the whole prompt, so small wording changes can change "
                "the answer a lot. (1 mark)",
                "So prompt variations must be TESTED, not assumed — the same model gives "
                "different-quality output for different prompts. (1 mark)",
            ]),
            ("(b) An example + two principles — 4 marks", [
                "Example from the activity: 'analyse this excel' gave a vague answer, while 'As a "
                "marketing analyst, list the top 3 channels by ROI with one action each' gave a "
                "usable one. (2 marks for a concrete before/after from their own work.)",
                "Principle 1: give the AI a ROLE. (1 mark)",
                "Principle 2 (any one): state the TASK exactly; give the CONTEXT/data; specify "
                "the FORMAT; set CONSTRAINTS (length, tone); give an example; ask for steps. "
                "(1 mark)",
            ]),
        ],
        guidance="Accept any genuine example from the learner's Activity 2-4 work. Two valid "
                 "principles earn the marks. 'Be more specific' with no principle named caps (b) "
                 "at 2 of 4.",
    ),
    5: dict(
        source="Topic 1 & 3 — How a model works; the agent as a system; AI-agent cybersecurity "
               "risks and safe rollout; Activities 7-8 (chatbot security).",
        points=[
            ("(a) How the model works — 4 marks", [
                "Training data: a very large body of text (and other media) is used to train the "
                "model; its knowledge and its biases come from that data. (1 mark)",
                "Algorithm at inference: it predicts the next token by probability using the "
                "context, one token at a time. (1 mark)",
                "Outputs: fluent, confident text that is NOT automatically fact-checked — "
                "plausible, not guaranteed correct, and it can vary between runs. (1 mark)",
                "Training happens once and is costly; inference is every time you prompt it. "
                "(1 mark)",
            ]),
            ("(b) Why an agent is riskier + two controls — 4 marks", [
                "As an agent, a wrong ANSWER becomes a wrong ACTION on a real system — the output "
                "boundary disappears. (1 mark)",
                "The loop chains actions with no human in the path, so one bad step can lead to "
                "another; some actions are hard to undo (for example, the Replit agent deleting a "
                "database). (1 mark)",
                "Control 1 (any one): least privilege — give the agent only the data and tools it "
                "needs, separating read from write. (1 mark)",
                "Control 2 (any one): a human approval step before risky or irreversible actions; "
                "sandboxing; egress/recipient limits; monitoring and a kill switch; testing "
                "against attacks before go-live. (1 mark)",
            ]),
        ],
        guidance="Part (a) is about the MODEL — a candidate who jumps straight to agents caps at "
                 "1 mark for (a). In (b) the two controls must differ in kind; a prompt-based "
                 "'tell the agent to be careful' earns 0 — a prompt is not a control.",
    ),
}

WA_INSTRUCTIONS = [
    "1. This is an INDIVIDUAL, OPEN-BOOK written assessment of underpinning knowledge.",
    "2. Answer ALL FIVE (5) questions. Every question is open-ended — there are no "
    "multiple-choice options.",
    "3. A total of 35 minutes is allowed for this Written Assessment.",
    "4. Write your answers in your own words in the answer box beneath each question. Continue on "
    "the reverse of the page if you need more space, labelling the question number clearly.",
    "5. Each question carries 8 marks (total 40 marks). The marks for each part are shown "
    "in brackets.",
    "6. Each question is tagged with the knowledge statement it assesses (e.g. [K2]). All "
    "questions are drawn from content taught across Topics 1 to 3.",
    "7. You must achieve a satisfactory response to ALL five questions to be assessed as "
    "Competent.",
    "8. Where you name a framework or a concept (for example the agentic loop, IMDA, PDPA), "
    "name it correctly — an invented identifier will not be credited.",
]


# ================================================================== PRACTICAL content
# The Practical Performance is a reflection on the activities the learner completed
# during the day. There is one task per Learning Outcome; together they cover all
# five accredited Ability statements. No fictional company or data tables are used —
# the "evidence" is the learner's own observations from the in-class activities.
CS_ORG = "Your own activity work"

CS_SCENARIO_INTRO = (
    "This Practical Performance is based on the eight hands-on activities you completed today. "
    "You will NOT analyse a fictional company. Instead, you will document your own observations "
    "and reflections from the activities you actually did, using the ready-made websites, "
    "chatbots and AI agents provided in class. Answer from your own notes, the Pinboard posts "
    "you made, and what you saw on screen."
)

CS_SCENARIO_BODY = [
    ("The activities you completed today", None,
     "Use these as the evidence for your answers. Refer to what you actually observed."),

    (None, "bullets", [
        "Activity 1 — you messaged the TIA Support AI agent on WhatsApp (powered by OpenClaw), "
        "gave it a task, and posted reflections to the Pinboard under Data Privacy, Job Impact, "
        "Ethical Concerns and Cyber Security.",
        "Activities 2-4 — you sent prompts to the Hermes agent (on a MiniMax model) to analyse "
        "mock Excel marketing data and to build an animated PowerPoint, comparing weak and strong "
        "prompts, then added a tool/skill and re-ran the tasks.",
        "Activity 5 — you adapted a sample AI Data Governance Policy to a team of your choice.",
        "Activity 6 — you coached an AI-played worried team member in the role-play simulator and "
        "received GROW-model feedback.",
        "Activities 7-8 — you made a deliberately insecure chatbot leak fictional PII, compared a "
        "guarded chatbot that blocked the same attacks, and decided whether an agent like it "
        "should go live.",
    ]),

    ("How to answer", None,
     "Write from your own experience. Where a task asks for an example, use something you "
     "actually saw or produced today. All the data you used was fictional, so you may quote it "
     "freely. There are no data tables to read — your activity notes are your evidence."),
]

# The Practical uses no data tables. These are kept empty so the renderer stays happy.
CS_TABLE_A = None
CS_TABLE_A2 = None
CS_TABLE_B = None
CS_TABLE_SEG = None

CS_PRESSURES = []

CS_ROLE = (
    "You are reflecting as a service or hospitality professional who has just trialled these AI "
    "tools. Answer all three tasks below from your own activity work. Each task maps to one "
    "Learning Outcome."
)

# (qno, LO, [A codes], marks, stem, [sub-prompts], answer_lines)
CS_Q = [
    (1, "LO1", ["A3", "A4"], 20,
     "Reflect on Activities 1-4, where you used generative AI and an AI agent to do real service "
     "tasks (summarising, drafting, analysing data and building slides).",
     ["(a) Describe ONE task you gave a generative AI tool or agent today and what it produced. "
      "State which kind of application it was — for example summarisation, drafting, analysis or "
      "transformation of content. (8 marks)",
      "(b) Using your Activity 1 experience with the AI agent on WhatsApp, explain in your own "
      "words the difference between generative AI, agentic AI and an AI agent. (8 marks)",
      "(c) Give ONE example from today where the AI was useful and ONE where you would not fully "
      "trust it, and say why. (4 marks)"],
     22),

    (2, "LO2", ["A5"], 24,
     "Reflect on Activities 2-4, where you compared weak and strong prompts on the same data and "
     "then added a tool or skill.",
     ["(a) Write out one WEAK prompt and one STRONG prompt you used today for the same task, and "
      "describe how the two outputs differed. (10 marks)",
      "(b) Explain which prompt-engineering principles made the stronger prompt work (for "
      "example role, task, context, format, constraints). (8 marks)",
      "(c) Describe what changed when you added a tool or skill in Activity 4, and what this "
      "tells you about improving AI performance beyond the prompt. (6 marks)"],
     22),

    (3, "LO3", ["A1", "A2"], 26,
     "Reflect on Activities 5-8, covering data governance, job impact and chatbot security.",
     ["(a) From Activity 7, describe what the leaky chatbot exposed and how the guarded chatbot "
      "stopped the same attack. Name at least two guardrail layers you saw. (8 marks)",
      "(b) Identify ONE ethical risk or bias you noticed in AI-generated content today, and its "
      "possible impact on a customer or member of the public. (8 marks)",
      "(c) From your Activity 5 policy and your Activity 6 coaching, explain who should be "
      "accountable when an AI agent handles data or affects a person's job, and why it is never "
      "the AI itself. (6 marks)",
      "(d) Using the safe-rollout idea from Activity 8, give a go, conditional-go or no-go "
      "decision for putting an agent like today's chatbot in front of real customers, and name "
      "one condition or owner. (4 marks)"],
     22),
]

CS_ANSWERS = {
    1: dict(
        source="Topic 1 — generative AI, agentic AI and AI agents; Activities 1-4.",
        blocks=[
            ("(a) A generative task and its type — 8 marks", [
                "Full marks for a clear, specific task the candidate actually did today, with the "
                "output described and the correct application type named.",
                "Accept any of: summarisation (condensing complaints or a document), drafting a "
                "reply, inference (likely cause of a delay), reasoning, transformation (reformat "
                "data or build slides), augmentation (answering from supplied documents).",
                "Award part marks if the task is described but the type is not named or is "
                "slightly off.",
            ]),
            ("(b) Generative vs agentic vs agent — 8 marks", [
                "Generative AI = makes content (a step you run). (about 3 marks)",
                "Agentic AI = the loop that makes the AI pursue a goal — plan, act, verify, "
                "repeat. (about 3 marks)",
                "AI agent = the deployed system that uses that loop to act for you, e.g. the "
                "WhatsApp agent that received the task and acted. (about 2 marks)",
                "Accept the summary 'generation is a step, agency is a loop, an agent is a "
                "system'.",
            ]),
            ("(c) Useful vs not-trusted example — 4 marks", [
                "One concrete 'useful' example (e.g. a good draft reply) — 2 marks.",
                "One concrete 'would not trust' example with a reason (e.g. it stated a fact it "
                "could not know; it could see personal data) — 2 marks.",
            ]),
        ],
        guidance="This is a reflection task — reward honest, specific observations from the "
                 "candidate's OWN activity work over textbook definitions. A generic answer with "
                 "no reference to what they did today caps at half marks. Satisfactory when parts "
                 "(a) and (b) are both addressed with a real example.",
    ),
    2: dict(
        source="Topic 2 — prompt engineering and output variation; Activities 2-4.",
        blocks=[
            ("(a) Weak vs strong prompt and the difference — 10 marks", [
                "A real weak prompt and a real strong prompt for the SAME task — 4 marks.",
                "A clear description of how the outputs differed (the strong one was usable, "
                "structured, on-format; the weak one vague) — 6 marks.",
                "Example: weak 'analyse this excel' vs strong 'As a marketing analyst, list the "
                "top 3 channels by ROI with one action each'.",
            ]),
            ("(b) Principles that made it work — 8 marks", [
                "Credit any principles the candidate correctly links to their strong prompt: "
                "role, task, context, format, constraints, giving an example, asking for steps. "
                "About 2 marks each, up to 8.",
            ]),
            ("(c) Effect of adding a tool/skill — 6 marks", [
                "A clear before/after from Activity 4 (better structure, formatting or accuracy) "
                "— 4 marks.",
                "The insight that performance can be improved beyond the prompt, by giving the "
                "agent tools/skills — 2 marks.",
            ]),
        ],
        guidance="Reward a genuine before/after from the candidate's own prompts. If they cannot "
                 "produce two contrasting prompts they did today, (a) caps at 4. Satisfactory "
                 "when (a) and (b) show real understanding of why the stronger prompt worked.",
    ),
    3: dict(
        source="Topic 3 — data governance, job impact, AI-agent security; Activities 5-8.",
        blocks=[
            ("(a) Leaky vs guarded chatbot — 8 marks", [
                "What the leaky bot exposed (fictional customer PII, internal memo, admin "
                "password) — 4 marks.",
                "How the guarded bot stopped it, naming at least two layers: retrieval filter, "
                "input guard, hardened prompt, output guard (redaction) — 4 marks.",
            ]),
            ("(b) An ethical risk or bias + impact — 8 marks", [
                "A specific ethical risk or bias noticed today (e.g. a confident but wrong "
                "answer; a leak of personal data; unfair treatment of one group) — 4 marks.",
                "Its possible impact on a customer or member of the public — 4 marks.",
            ]),
            ("(c) Accountability — 6 marks", [
                "A named human role is accountable (data owner, agent owner, manager) — 3 marks.",
                "Why it is never the AI: the AI is a tool with no legal responsibility; the "
                "organisation answers for it (accept the Moffatt v Air Canada point) — 3 marks.",
            ]),
            ("(d) Go / no-go decision — 4 marks", [
                "A clear go, conditional-go or no-go decision — 2 marks.",
                "At least one condition, owner or kill-switch named — 2 marks.",
            ]),
        ],
        guidance="Reward reflection grounded in what the candidate saw in Activities 5-8. The "
                 "accountability point (never the AI) must appear for full marks in (c). "
                 "Satisfactory when (a), (b) and (c) are addressed with real observations.",
    ),
}

CS_INSTRUCTIONS = [
    "1. This is an INDIVIDUAL, OPEN-BOOK practical reflection assessment.",
    "2. Your answers are based on the eight activities you completed during the course today.",
    "3. Answer ALL THREE (3) tasks. Every task is open-ended — there are no multiple-choice "
    "options.",
    "4. A total of 25 minutes is allowed for this Practical Performance.",
    "5. Write from your own experience. Where a task asks for an example, use something you "
    "actually did or saw today. All the data you used was fictional, so you may quote it freely.",
    "6. Marks: Task 1 = 20 marks, Task 2 = 24 marks, Task 3 = 26 marks (total 70 marks). The "
    "marks for each part are shown in brackets.",
    "7. Each task is tagged with the Learning Outcome and the ability statement(s) it assesses "
    "(e.g. [LO1 · A4]).",
    "8. You must achieve a satisfactory response to ALL three tasks to be assessed as Competent.",
]


# ================================================================== table renderer
def render_table(d, spec, font=8.5):
    line(d, spec["caption"], bold=True, size=9.5, color=BRAND, after=4, before=4)
    t = d.add_table(rows=1, cols=len(spec["head"]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(spec["head"]):
        _cell_text(t.rows[0].cells[i], h, bold=True, size=font, color=WHITE, align=AL.CENTER)
        _shade(t.rows[0].cells[i], "1F6FEB")
    for r_i, row in enumerate(spec["rows"]):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            _cell_text(cells[i], v, size=font,
                       align=AL.CENTER if i > 0 and len(str(v)) < 12 else None)
        if r_i % 2 == 1:
            for c in cells:
                _shade(c, "F5F7FB")
    if spec.get("widths"):
        _fix_widths(t, spec["widths"])
    line(d, "", after=6)


# ================================================================== coverage table
def coverage_table(d, kind):
    line(d, "ASSESSMENT COVERAGE MAPPING", bold=True, size=13, color=BRAND, after=3)
    if kind == "WA":
        line(d, "Every accredited Knowledge (K) statement of TSC " + TSC_CODE + " is assessed by "
                "exactly one question in this Written Assessment. No K statement is unassessed "
                "and no question assesses a statement not listed below.",
             size=10, color=GREY, after=8)
        head = ["Question", "K Code", "Accredited Knowledge Statement (verbatim)", "Topic", "Marks"]
        rows = [
            ["Q1", "K2", K_STATEMENTS["K2"], "T1", "8"],
            ["Q2", "K3", K_STATEMENTS["K3"], "T1", "8"],
            ["Q3", "K1", K_STATEMENTS["K1"], "T3", "8"],
            ["Q4", "K4", K_STATEMENTS["K4"], "T2", "8"],
            ["Q5", "K5", K_STATEMENTS["K5"], "T1/T3", "8"],
        ]
        widths = [0.75, 0.7, 3.5, 0.7, 0.65]
    else:
        line(d, "Every accredited Ability (A) statement of TSC " + TSC_CODE + " is assessed "
                "across the three Practical tasks, one task per Learning Outcome. All five A "
                "statements are covered; none is unassessed.",
             size=10, color=GREY, after=8)
        head = ["Task", "LO", "A Code", "Accredited Ability Statement (verbatim)", "Marks"]
        rows = [
            ["T1", "LO1", "A3", A_STATEMENTS["A3"], "20"],
            ["T1", "LO1", "A4", A_STATEMENTS["A4"], "(within T1)"],
            ["T2", "LO2", "A5", A_STATEMENTS["A5"], "24"],
            ["T3", "LO3", "A1", A_STATEMENTS["A1"], "26"],
            ["T3", "LO3", "A2", A_STATEMENTS["A2"], "(within T3)"],
        ]
        widths = [0.72, 0.55, 0.62, 3.75, 0.86]

    t = d.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(head):
        _cell_text(t.rows[0].cells[i], h, bold=True, size=8.5, color=WHITE, align=AL.CENTER)
        _shade(t.rows[0].cells[i], "1F6FEB")
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            _cell_text(cells[i], v, size=8.5,
                       align=AL.CENTER if i != len(row) - 2 else None)
        if r_i % 2 == 1:
            for c in cells:
                _shade(c, "F5F7FB")
    _fix_widths(t, widths)
    line(d, "", after=4)

    if kind == "WA":
        line(d, "Coverage confirmation: K1 ✓ (Q3) · K2 ✓ (Q1) · K3 ✓ (Q2) · K4 ✓ (Q4) · "
                "K5 ✓ (Q5) — 5 of 5 knowledge statements assessed.",
             bold=True, size=10, color=DARK, after=4)
    else:
        line(d, "Coverage confirmation: A1 ✓ (T3c) · A2 ✓ (T3a, T3b) · A3 ✓ (T1a, T1b) · "
                "A4 ✓ (T1a–c) · A5 ✓ (T2a–c) — 5 of 5 ability statements assessed across "
                "3 tasks, one per Learning Outcome.",
             bold=True, size=10, color=DARK, after=4)
    keep_with_next(line(d, "Learning Outcome mapping:", bold=True, size=10, after=3))
    if kind == "WA":
        bullet(d, "LO1 (Topic 1) — Q1 [K2], Q2 [K3]", size=9.5)
        bullet(d, "LO2 (Topic 2) — Q4 [K4]", size=9.5)
        bullet(d, "LO3 (Topic 3) — Q3 [K1], Q5 [K5]", size=9.5)
    else:
        for lo in ("LO1", "LO2", "LO3"):
            keep_lines(bullet(d, f"{lo} — {LO_STATEMENTS[lo]}", size=9.5))


# ================================================================== builders
def build_wa_paper():
    d = new_doc()
    footer_block(d)
    enable_update_fields(d)
    cover_page(d, "Written Assessment (SAQ)", "Question Paper")           # page 1
    admin_page(d, "Written Assessment (SAQ)", "35 minutes", 40, 5,
               "short-answer questions", WA_INSTRUCTIONS)                 # page 2

    # ------------------------------------------------ page 3 onwards
    line(d, "SECTION D: SHORT-ANSWER QUESTIONS", bold=True, size=13, color=BRAND, after=3)
    line(d, "Answer ALL FIVE questions. Each question carries 8 marks. Total: 40 marks.",
         bold=True, size=10.5, color=DARK, after=10)

    for idx, (n, k, marks, ctx, stem, subs, nlines) in enumerate(WRITTEN):
        runs(d, [(f"Question {n}  ", True), (f"[{k}]", True),
                 (f"   ({marks} marks)", False)], after=3, size=12)
        line(d, f"Knowledge statement {k}: {K_STATEMENTS[k]}", size=8.5, color=GREY,
             italic=True, after=3)
        line(d, f"Context — {ctx}", size=9, color=GREY, italic=True, after=5)
        line(d, stem, size=11, after=6)
        for s in subs:
            line(d, s, size=11, after=5, indent=0.2)
        answer_box(d, nlines)
        if idx < len(WRITTEN) - 1:
            page_break(d)

    out = os.path.join(
        OUT, f"WSQ - Written Assessment (SAQ) - Question Paper - {COURSE_CODE} - {TITLE}.docx")
    d.save(out)
    return out


def build_wa_answers():
    d = new_doc()
    footer_block(d)
    enable_update_fields(d)
    cover_page(d, "Written Assessment (SAQ)", "Answer Key")               # page 1

    # ------------------------------------------------ page 2: assessor note + coverage
    line(d, "Written Assessment (SAQ) — Answer Key and Marking Guide", bold=True, size=15,
         color=BRAND, after=2, align=AL.CENTER)
    line(d, TITLE, bold=True, size=12, color=DARK, after=2, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}  ·  TSC: {TSC_CODE}  ·  Version {VERSION}  ·  "
            f"{VERSION_DATE}", size=9.5, color=GREY, after=12, align=AL.CENTER)

    line(d, "NOTE TO ASSESSOR", bold=True, size=12, color=BRAND, after=4)
    for t in [
        "This paper contains FIVE open-ended short-answer questions, one for each accredited "
        "Knowledge statement of TSC " + TSC_CODE + ". Each question is worth 8 marks; the paper "
        "totals 40 marks.",
        "The model answers below are NOT a script. Award the mark wherever the candidate covers "
        "the underlying point — wording, ordering and terminology will vary, and a candidate may "
        "reach a correct conclusion by a route not shown here. Acceptable alternatives are "
        "listed in square brackets and in the marking guidance under each question.",
        "A question is SATISFACTORY when the candidate scores 5 or more of its 8 marks AND "
        "addresses both parts. The candidate is COMPETENT only when all five questions are "
        "satisfactory (and therefore at least 25 of 40 overall).",
        "OWASP identifiers (LLM01–LLM10, ASI01–ASI10) are a bonus, not a requirement — a "
        "candidate who explains the mechanism correctly in plain language earns full marks. "
        "However, an INVENTED identifier is a defect: do not credit it, and note it in the "
        "feedback.",
        "This is an open-book assessment. Verbatim reproduction of slide text without "
        "application to the question stem does not evidence understanding — probe with an oral "
        "question where an answer looks copied.",
    ]:
        bullet(d, t, size=11)
    line(d, "", after=10)
    coverage_table(d, "WA")
    page_break(d)

    # ------------------------------------------------ one model answer per page
    for idx, (n, k, marks, ctx, stem, subs, _) in enumerate(WRITTEN):
        a = WA_ANSWERS[n]
        keep_with_next(runs(d, [(f"Question {n}  ", True), (f"[{k}]", True),
                                (f"   ({marks} marks)", False)], after=3, size=13))
        line(d, f"Knowledge statement {k}: {K_STATEMENTS[k]}", size=9, color=GREY,
             italic=True, after=3)
        line(d, f"Taught in: {a['source']}", size=9, color=GREY, italic=True, after=6)
        line(d, stem, size=10.5, color=GREY, after=4)
        for s in subs:
            line(d, s, size=10.5, color=GREY, after=3, indent=0.2)
        line(d, "", after=4)

        keep_with_next(line(d, "Model answer (not exhaustive):", bold=True, size=11,
                            color=DARK, after=4))
        for head, pts in a["points"]:
            line(d, head, bold=True, size=10.5, color=BRAND, after=3, before=2)
            for p in pts:
                bullet(d, p, size=11)
        line(d, "", after=4)
        keep_with_next(line(d, "Marking guidance and acceptable alternatives:", bold=True,
                            size=11, color=DARK, after=3))
        line(d, a["guidance"], size=11, after=14)

    out = os.path.join(
        OUT, f"WSQ - Written Assessment (SAQ) - Answer Key - {COURSE_CODE} - {TITLE}.docx")
    d.save(out)
    return out


def scenario_pages(d, with_role=True):
    """Renders the Practical reflection context. Starts on the CURRENT page."""
    line(d, "SECTION D: YOUR ACTIVITY WORK TODAY", bold=True, size=13, color=BRAND, after=3)
    line(d, "Practical Performance — reflection on the activities you completed",
         bold=True, size=12, color=DARK, after=8)
    line(d, CS_SCENARIO_INTRO, size=11, after=8)

    for heading, mode, body in CS_SCENARIO_BODY:
        if mode == "bullets":
            for b in body:
                bullet(d, b, size=11)
            line(d, "", after=4)
        else:
            if heading:
                line(d, heading, bold=True, size=11, color=BRAND, after=3, before=4)
            line(d, body, size=11, after=5)

    # Optional data tables — only rendered if present (the reflection Practical has none).
    for tbl in (CS_TABLE_A, CS_TABLE_A2, CS_TABLE_B, CS_TABLE_SEG):
        if tbl:
            render_table(d, tbl)

    for i, p in enumerate(CS_PRESSURES, 1):
        runs(d, [(f"{i}. ", True), (p, False)], after=5, size=11)

    if with_role:
        line(d, "Your role", bold=True, size=12, color=BRAND, after=4, before=6)
        line(d, CS_ROLE, size=11, after=6)


def build_cs_paper():
    d = new_doc()
    footer_block(d)
    enable_update_fields(d)
    cover_page(d, "Practical Performance", "Question Paper")              # page 1
    admin_page(d, "Practical Performance", "25 minutes", 70, 3, "tasks",
               CS_INSTRUCTIONS)                                           # page 2

    scenario_pages(d)                                                     # page 3 onwards
    page_break(d)

    line(d, "SECTION E: PRACTICAL REFLECTION TASKS", bold=True, size=13, color=BRAND, after=3)
    line(d, "Answer ALL THREE tasks. Total: 70 marks.", bold=True, size=10.5,
         color=DARK, after=10)

    # Each task has a multi-part stem. Putting the stem and a full-height answer box on
    # one page overflows and fragments the box, so the stem gets its own page and the
    # answer box starts on a fresh page at full height.
    for idx, (n, lo, acodes, marks, stem, subs, nlines) in enumerate(CS_Q):
        tag = f"[{lo} · {' · '.join(acodes)}]"
        runs(d, [(f"Task {n}  ", True), (tag, True),
                 (f"   ({marks} marks)", False)], after=3, size=12)
        line(d, f"Learning Outcome {lo}: {LO_STATEMENTS[lo]}", size=8.5, color=GREY,
             italic=True, after=2)
        for ac in acodes:
            line(d, f"Ability statement {ac}: {A_STATEMENTS[ac]}", size=8.5, color=GREY,
                 italic=True, after=2)
        line(d, "", after=3)
        line(d, stem, size=11, after=6)
        for s in subs:
            line(d, s, size=11, after=6, indent=0.2)

        page_break(d)
        line(d, f"Task {n} — Answer", bold=True, size=11, color=BRAND, after=4)
        answer_box(d, nlines)
        if idx < len(CS_Q) - 1:
            page_break(d)

    out = os.path.join(
        OUT, f"WSQ - Case Study - Question Paper - {COURSE_CODE} - {TITLE}.docx")
    d.save(out)
    return out


def build_cs_answers():
    d = new_doc()
    footer_block(d)
    enable_update_fields(d)
    cover_page(d, "Practical Performance", "Answer Key")                  # page 1

    line(d, "Practical Performance — Answer Key and Marking Guide", bold=True, size=15,
         color=BRAND, after=2, align=AL.CENTER)
    line(d, TITLE, bold=True, size=12, color=DARK, after=2, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}  ·  TSC: {TSC_CODE}  ·  Version {VERSION}  ·  "
            f"{VERSION_DATE}", size=9.5, color=GREY, after=12, align=AL.CENTER)

    line(d, "NOTE TO ASSESSOR", bold=True, size=12, color=BRAND, after=4)
    for t in [
        "This is a REFLECTION-BASED practical. The candidate answers THREE open-ended tasks about "
        "the eight activities they completed during the course, one task per Learning Outcome. "
        "Together the three tasks cover all five accredited Ability statements of TSC " +
        TSC_CODE + ".",
        "Marks: Task 1 = 20, Task 2 = 24, Task 3 = 26. Total 70 marks.",
        "Because answers are drawn from each learner's OWN activity work, they will differ. Award "
        "the mark wherever the candidate demonstrates the underlying point with a genuine example "
        "from today's activities. Reward specific, honest reflection over textbook wording.",
        "Each task states its own satisfactory threshold in the marking guidance. The candidate "
        "is COMPETENT only when ALL THREE tasks are satisfactory.",
        "The activity context reproduced below is identical to the candidate's paper.",
        "Where a candidate names a concept or framework (the agentic loop, a guardrail layer, "
        "IMDA, PDPA), it must be correct — but the emphasis is on their observed experience, not "
        "on reciting identifiers.",
    ]:
        bullet(d, t, size=11)
    line(d, "", after=10)
    coverage_table(d, "CS")
    page_break(d)

    scenario_pages(d, with_role=False)
    page_break(d)

    for idx, (n, lo, acodes, marks, stem, subs, _) in enumerate(CS_Q):
        a = CS_ANSWERS[n]
        tag = f"[{lo} · {' · '.join(acodes)}]"
        keep_with_next(runs(d, [(f"Task {n}  ", True), (tag, True),
                                (f"   ({marks} marks)", False)], after=3, size=13))
        line(d, f"Learning Outcome {lo}: {LO_STATEMENTS[lo]}", size=9, color=GREY,
             italic=True, after=2)
        for ac in acodes:
            line(d, f"Ability statement {ac}: {A_STATEMENTS[ac]}", size=9, color=GREY,
                 italic=True, after=2)
        line(d, f"Taught in: {a['source']}", size=9, color=GREY, italic=True, after=6)
        line(d, stem, size=10.5, color=GREY, after=4)
        for s in subs:
            line(d, s, size=10.5, color=GREY, after=3, indent=0.2)
        line(d, "", after=4)

        keep_with_next(line(d, "Model answer (not exhaustive):", bold=True, size=11,
                            color=DARK, after=4))
        for head, pts in a["blocks"]:
            line(d, head, bold=True, size=10.5, color=BRAND, after=3, before=3)
            for p in pts:
                bullet(d, p, size=11)
        line(d, "", after=4)
        keep_with_next(line(d, "Marking guidance and acceptable alternatives:", bold=True,
                            size=11, color=DARK, after=3))
        line(d, a["guidance"], size=11, after=14)

    out = os.path.join(
        OUT, f"WSQ - Case Study - Answer Key - {COURSE_CODE} - {TITLE}.docx")
    d.save(out)
    return out


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    for fn in (build_wa_paper, build_wa_answers, build_cs_paper, build_cs_answers):
        print("Wrote:", os.path.basename(fn()))
