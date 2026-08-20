#!/usr/bin/env python3
"""WSQ Learner Guide — AI Security for Autonomous AI Agents (TGS-2025060473).

Emits BOTH the DOCX and an aligned Markdown mirror from lg_content.SECTIONS,
so the two can never diverge.
"""
import os, sys

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL = os.path.join(HERE, "..", ".claude", "skills", "tertiary-lesson-plan")
sys.path.insert(0, os.path.abspath(SKILL))
sys.path.insert(0, HERE)

import prodoc
prodoc.ORG = "Tertiary Infotech Pte Ltd"
prodoc.UEN = "UEN: 20120096W"
prodoc.COPYRIGHT = ("This material belongs to Tertiary Infotech Pte Ltd (UEN: 20120096W). "
                    "All Rights Reserved.")

from prodoc import (add_cover_page, add_version_control, add_toc, add_page_numbers,
                    enable_update_fields, style_headings, _shade_cell)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import course_data as C
if C.VERSION.startswith("3"):
    from v30_learner import SECTIONS
else:
    from lg_content import SECTIONS

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66); TEAL = RGBColor(0x10, 0xB9, 0x81)
ASSETS = os.path.join(HERE, "assets")

VERSIONS = [
    ("1.0", "16 June 2025", "Initial release — Core Principles and Ethical Challenges in "
                            "Generative AI.", "Dr Alfred Ang"),
    ("2.0", "17 August 2026", "Retitled and rebuilt for generative-AI and autonomous-agent security; "
                            "five activities and detailed walkthroughs added. TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("2.1", "18 August 2026", "Added prompt-injection practice, PDPA guidance, guardrails, human "
                            "approval, skill/plugin controls and visual teaching aids. TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("3.0", "20 August 2026", "Aligned to the evidence-grounded 207-slide deck; added product "
                            "boundaries, documented cases, control gates, detailed activities and "
                            "source register. Unverified claims removed. TSC K/A text unchanged.",
     "Dr Alfred Ang"),
]


def para(doc, t, size=11, bold=False, color=DARK, after=8, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    r.font.name = "Arial"; r.font.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p


def table(doc, headers, rows, font_size=9, keep_whole=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.font.bold = True; r.font.size = Pt(font_size + 0.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = "Arial"
        _shade_cell(c, "1F6FEB")
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font_size); r.font.name = "Arial"; r.font.bold = (j == 0)
            r.font.color.rgb = DARK
            cells[j].paragraphs[0].paragraph_format.space_after = Pt(0)
        if i % 2 == 0:
            for c in cells:
                _shade_cell(c, "F5F8FC")
    # Keep short evidence tables with their headings and prevent individual rows
    # from splitting. The long source register may paginate between complete rows.
    for row_i, row in enumerate(t.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                if keep_whole:
                    p.paragraph_format.keep_with_next = row_i < len(t.rows) - 1
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    p1 = c.paragraphs[0]; r1 = p1.add_run(title.upper())
    r1.font.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = BRAND; r1.font.name = "Arial"
    p2 = c.add_paragraph(); r2 = p2.add_run(body)
    r2.font.size = Pt(10); r2.font.color.rgb = DARK; r2.font.name = "Arial"
    _shade_cell(c, "EEF3FB")
    keep_next = title.upper().startswith("EVIDENCE STATUS")
    for p in c.paragraphs:
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = keep_next
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build_docx():
    doc = Document()
    style_headings(doc)
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].paragraph_format.keep_with_next = True
        doc.styles[name].paragraph_format.keep_together = True

    add_cover_page(doc, "Learner Guide", C.TITLE, C.VERSION,
                   conducted_by="Tertiary Infotech Pte Ltd",
                   org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                   course_logo=None, course_code=C.COURSE_CODE)

    doc.add_heading("Document Version Control Record", level=1)
    add_version_control(doc, VERSIONS)

    doc.add_heading("Table of Contents", level=1)
    para(doc, "Generated section contents for this v3.0 build.", size=9.5, color=GREY)
    table(doc, ["Section", "Page"], [
        ["About This Course", "4"],
        ["How to Use This Evidence-Grounded Guide", "5"],
        ["Day 1 — From AI Models to Acting Systems", "9"],
        ["Day 2 — Attack, Defend, Govern, Operate", "50"],
        ["Detailed Activity Walkthroughs", "101"],
        ["Operational Best-Practice Checklist", "107"],
        ["Source Register", "108"],
    ])
    doc.add_page_break()

    # front matter
    doc.add_heading("About This Course", level=1)
    para(doc, f"{C.TITLE} is a {C.DURATION} WSQ course covering the security of generative AI "
              "systems and of autonomous AI agents. It is delivered through real-world case "
              "studies and grounded in the current security and governance frameworks.")
    para(doc, "This guide follows the course structure. Each section names the accredited "
              "knowledge (K) or ability (A) statement it evidences, and every activity has a "
              "full step-by-step walkthrough.")

    doc.add_heading("Learning Outcomes", level=2)
    table(doc, ["LO", "Learning Outcome"], [[c, d] for c, d in C.LEARNING_OUTCOMES])

    doc.add_heading("Knowledge Statements", level=2)
    table(doc, ["Code", "Knowledge statement"], [[c, d] for c, d in C.TSC_KNOWLEDGE])

    doc.add_heading("Ability Statements", level=2)
    table(doc, ["Code", "Ability statement"], [[c, d] for c, d in C.TSC_ABILITIES])
    # Body. Do not force an additional break before the first H1; the front-matter
    # tables may already fill the page and an explicit break can create a blank page.
    def _drop_trailing_empty_paragraphs():
        """table()/callout() end with a spacer paragraph. If one is the last element
        before a page break, the page overflows and the break lands on a fresh page,
        producing a blank one. Remove any trailing empties first."""
        body = doc.element.body
        while len(body) > 1:
            last = body[-1]
            if last.tag.endswith('}sectPr'):
                last = body[-2] if len(body) > 1 else None
                if last is None:
                    return
            if last.tag.endswith('}p') and not ''.join(last.itertext()).strip():
                body.remove(last)
            else:
                return

    first_h1 = True
    in_activity_walkthroughs = False
    slide_blocks_on_page = 0
    for level, kind, payload in SECTIONS:
        if kind == "h1":
            if not first_h1:
                _drop_trailing_empty_paragraphs()
                doc.add_page_break()
            first_h1 = False
            doc.add_heading(payload, level=1)
            in_activity_walkthroughs = payload == "Detailed Activity Walkthroughs"
            slide_blocks_on_page = 0
        elif kind == "h2":
            if in_activity_walkthroughs and payload.startswith("Activity "):
                _drop_trailing_empty_paragraphs()
                doc.add_page_break()
            elif payload.startswith("Slide "):
                # Two complete slide-learning blocks per page gives the learner
                # enough room for evidence, tables and control implications without
                # leaving headings or list fragments on the next page.
                if slide_blocks_on_page >= 2:
                    _drop_trailing_empty_paragraphs()
                    doc.add_page_break()
                    slide_blocks_on_page = 0
                slide_blocks_on_page += 1
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "p":
            para(doc, payload)
        elif kind == "bullets":
            compact = in_activity_walkthroughs
            for i, b in enumerate(payload):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(b); r.font.size = Pt(9.3 if compact else 10.5); r.font.name = "Arial"
                r.font.color.rgb = DARK
                p.paragraph_format.space_after = Pt(3 if compact else 5)
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = i < len(payload) - 1
        elif kind == "numbered":
            # Word's List Number style shares one continuous sequence across the whole
            # document, so each activity's steps would carry on from the previous one
            # (Activity 2 starting at 10, etc.). Number them literally instead so every
            # walkthrough restarts at 1.
            compact = in_activity_walkthroughs
            for i, b in enumerate(payload, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.first_line_indent = Inches(-0.4)
                rn = p.add_run(f"{i}.  ")
                rn.font.size = Pt(9.3 if compact else 10.5); rn.font.name = "Arial"; rn.font.bold = True
                rn.font.color.rgb = BRAND
                r = p.add_run(b); r.font.size = Pt(9.3 if compact else 10.5); r.font.name = "Arial"
                r.font.color.rgb = DARK
                p.paragraph_format.space_after = Pt(3 if compact else 5)
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = i < len(payload)
        elif kind == "table":
            is_source_register = payload[0] == ["ID", "Source", "URL"]
            table(doc, payload[0], payload[1],
                  font_size=7.2 if is_source_register else (8.4 if in_activity_walkthroughs else 9),
                  keep_whole=not is_source_register)
        elif kind == "callout":
            callout(doc, payload[0], payload[1])
        elif kind == "image":
            fn, cap = payload
            path = os.path.join(ASSETS, fn)
            if os.path.exists(path):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(6.1))
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(cap)
                cr.font.size = Pt(9); cr.font.italic = True
                cr.font.color.rgb = GREY; cr.font.name = "Arial"
                cp.paragraph_format.space_after = Pt(10)

    add_page_numbers(doc)
    enable_update_fields(doc)
    out = os.path.join(HERE, f"Learner Guide - {C.COURSE_CODE} - {C.TITLE}.docx")
    doc.save(out)
    return out


def build_md():
    L = []
    L.append(f"# {C.TITLE} — Learner Guide\n")
    L.append(f"**Course Code:** {C.COURSE_CODE}  |  **TSC:** {C.TSC_TITLE} ({C.TSC_CODE})  \n"
             f"**Version:** {C.VERSION}  |  **Date:** {C.VERSION_DATE}  |  **Duration:** {C.DURATION}\n")
    source_module = "v30_learner.py" if C.VERSION.startswith("3") else "lg_content.py"
    L.append("> This guide mirrors the Learner Guide DOCX exactly. Both are generated from "
             f"`{source_module}`.\n")

    L.append("## Learning Outcomes\n")
    L.append("| LO | Learning Outcome |")
    L.append("|---|---|")
    for c, d in C.LEARNING_OUTCOMES:
        L.append(f"| {c} | {d} |")
    L.append("")

    L.append("## Knowledge Statements\n")
    L.append("| Code | Knowledge statement |")
    L.append("|---|---|")
    for c, d in C.TSC_KNOWLEDGE:
        L.append(f"| {c} | {d} |")
    L.append("")

    L.append("## Ability Statements\n")
    L.append("| Code | Ability statement |")
    L.append("|---|---|")
    for c, d in C.TSC_ABILITIES:
        L.append(f"| {c} | {d} |")
    L.append("")

    for level, kind, payload in SECTIONS:
        if kind == "h1":
            L.append(f"\n---\n\n# {payload}\n")
        elif kind == "h2":
            L.append(f"\n## {payload}\n")
        elif kind == "h3":
            L.append(f"\n### {payload}\n")
        elif kind == "p":
            L.append(payload + "\n")
        elif kind == "bullets":
            for b in payload:
                L.append(f"- {b}")
            L.append("")
        elif kind == "numbered":
            for i, b in enumerate(payload, 1):
                L.append(f"{i}. {b}")
            L.append("")
        elif kind == "table":
            hdr, rows = payload
            L.append("| " + " | ".join(str(h) for h in hdr) + " |")
            L.append("|" + "---|" * len(hdr))
            for row in rows:
                L.append("| " + " | ".join(str(c).replace("\n", " ") for c in row) + " |")
            L.append("")
        elif kind == "callout":
            L.append(f"> **{payload[0]}**\n> {payload[1]}\n")
        elif kind == "image":
            fn, cap = payload
            L.append(f"![{cap}](assets/{fn})\n")
            L.append(f"*{cap}*\n")

    L.append("\n---\n")
    L.append("*This material belongs to Tertiary Infotech Pte Ltd (UEN: 20120096W). "
             "All Rights Reserved.*")

    out = os.path.join(HERE, "LEARNER-GUIDE.md")
    open(out, "w", encoding="utf-8").write("\n".join(L))
    return out


if __name__ == "__main__":
    d = build_docx(); print("Saved:", d)
    m = build_md(); print("Saved:", m)
