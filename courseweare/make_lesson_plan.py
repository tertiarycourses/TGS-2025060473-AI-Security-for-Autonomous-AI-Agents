#!/usr/bin/env python3
"""WSQ Lesson Plan — AI Security for Autonomous AI Agents (TGS-2025060473).

2 days × 8 instructional hours. 9:30am–6:30pm with a 1-hour lunch; tea breaks
counted within the day. Asserts 480 min/day before saving.
"""
import os, sys

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL = os.path.join(HERE, "..", ".claude", "skills", "tertiary-lesson-plan")
sys.path.insert(0, os.path.abspath(SKILL))
sys.path.insert(0, HERE)

import prodoc
# This course is delivered by Tertiary Infotech Pte Ltd (UEN 20120096W) — not the
# Academy entity hardcoded in the shared helper. Override before use.
prodoc.ORG = "Tertiary Infotech Pte Ltd"
prodoc.UEN = "UEN: 20120096W"
prodoc.COPYRIGHT = ("This material belongs to Tertiary Infotech Pte Ltd (UEN: 20120096W). "
                    "All Rights Reserved.")

from prodoc import (add_cover_page, add_version_control, add_toc, add_page_numbers,
                    enable_update_fields, style_headings, _shade_cell)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import course_data as C

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)
ASSETS = os.path.join(HERE, "assets")

VERSIONS = [
    ("1.0", "16 June 2025", "Initial release — Core Principles and Ethical Challenges in "
                            "Generative AI.", "Dr Alfred Ang"),
    ("2.0", "17 August 2026", "Retitled and rebuilt for generative-AI and autonomous-agent security; "
                            "five activities and revised assessments added. TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("2.1", "18 August 2026", "Added prompt-injection practice, PDPA and organisational checklists, "
                            "human approval, skill/plugin controls and enhanced visuals. TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("3.0", "20 August 2026", "Rebuilt as a 207-slide evidence-grounded progression from AI history "
                            "to product boundaries, cases, controls and governance. Activities "
                            "reordered; unverified claims removed. TSC K/A text unchanged.",
     "Dr Alfred Ang"),
]

# (time, topic, minutes, kind)  kind: topic | activity | break | admin | assess
DAY1 = [
    ("9:30 – 9:50", "Welcome, digital attendance, introductions, learning outcomes and safe-lab boundaries", 20, "admin"),
    ("9:50 – 10:35", "AI history: symbolic systems, machine learning, deep learning, transformers and instruction following (K2)", 45, "topic"),
    ("10:35 – 10:50", "Tea break", 15, "break"),
    ("10:50 – 11:35", "Generative AI, agentic behaviour and deployed AI agents as overlapping operating layers (K2, K3)", 45, "topic"),
    ("11:35 – 12:30", "Strengths, weaknesses and examples: OpenClaw, Hermes, Claude Code, Codex, ChatGPT Work, DeepSeek model/harness and evidence gates", 55, "topic"),
    ("12:30 – 1:30", "Lunch break", 60, "lunch"),
    ("1:30 – 2:15", "Capability, data, token, code-execution, skill/plugin/MCP, network and approval matrices", 45, "topic"),
    ("2:15 – 3:00", "Threat modelling: assets, actors, trust boundaries, authority paths and abuse cases (K1, K5)", 45, "topic"),
    ("3:00 – 3:45", "Activity 1: Threat Modelling a Generative AI Concierge (K2, K3, A4)", 45, "activity"),
    ("3:45 – 4:00", "Tea break", 15, "break"),
    ("4:00 – 4:45", "Attack surfaces by operating layer: content, state, identity, tools, runtime and network", 45, "topic"),
    ("4:45 – 5:30", "Complete attack chains, source-to-sink reasoning and evidence classification", 45, "topic"),
    ("5:30 – 6:15", "Integrated risk matrix and preview of structural controls, approvals and evidence", 45, "topic"),
    ("6:15 – 6:30", "Day 1 recap, Q&A and PM digital attendance", 15, "admin"),
]

DAY2 = [
    ("9:30 – 9:40", "Day 1 recap and AM digital attendance", 10, "admin"),
    ("9:40 – 10:00", "Prompt injection, jailbreaks, injection carriers and structural control principles (K4)", 20, "topic"),
    ("10:00 – 11:00", "Activity 2: Prompt Injection and the PDPA Breach Decision (K4, K1)", 60, "activity"),
    ("11:00 – 11:15", "Tea break", 15, "break"),
    ("11:15 – 11:40", "Personal data, secrets, tokens, retention, transfers and the PDPA breach decision", 25, "topic"),
    ("11:40 – 12:00", "Framework stack followed by detailed agentic-threat and documented-case instruction (A3, K5)", 20, "topic"),
    ("12:00 – 1:00", "Activity 3: Selecting a Security Framework for GenAI and Agents (A3, A5)", 60, "activity"),
    ("1:00 – 2:00", "Lunch break", 60, "lunch"),
    ("2:00 – 2:30", "Activity 3 debrief and evidence bridge into the documented incident review", 30, "topic"),
    ("2:30 – 3:30", "Activity 4: Evidence-Based Rogue Agent Incident Review (K5)", 60, "activity"),
    ("3:30 – 3:45", "Tea break", 15, "break"),
    ("3:45 – 4:10", "Organisational controls: guardrails, HITL, sandbox, egress, monitoring, deployment and incident gates (A1, A2)", 25, "topic"),
    ("4:10 – 4:35", "Activity 5: Agent Governance and the Deployment Gate (A1, A2)", 25, "activity"),
    ("4:35 – 4:50", "Course synthesis, assessment briefing and PM digital attendance", 15, "admin"),
    ("4:50 – 5:50", "Assessment 1: Written Assessment (SAQ) — 5 questions · 40 marks · 60 min", 60, "assess"),
    ("5:50 – 6:30", "Assessment 2: Case Study — 3 questions · 70 marks · 40 min; TRAQOM close", 40, "assess"),
]

SHADE = {"topic": "E8F0FE", "activity": "E6F7F1", "break": "FDF3E3", "lunch": "FDF3E3",
         "admin": "F2F4F7", "assess": "FDECEC"}


def h1(doc, t):
    p = doc.add_heading(t, level=1); return p


def h2(doc, t):
    return doc.add_heading(t, level=2)


def para(doc, t, size=11, bold=False, color=DARK, after=6):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    r.font.name = "Arial"
    p.paragraph_format.space_after = Pt(after)
    return p


def table(doc, headers, rows, widths=None, shades=None, keep_whole=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(htxt)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Arial"
        _shade_cell(c, "1F6FEB")
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5); r.font.name = "Arial"
            r.font.bold = (j == 0 and shades is None)
        if shades:
            for c in cells:
                _shade_cell(c, shades[i])
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    for row_i, row in enumerate(t.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                if keep_whole:
                    p.paragraph_format.keep_with_next = row_i < len(t.rows) - 1
    return t


def day_table(doc, day):
    rows = [(tm, tp, f"{mn} min") for tm, tp, mn, _ in day]
    shades = [SHADE[k] for _, _, _, k in day]
    table(doc, ["Time", "Topic / Activity", "Duration"], rows,
          widths=[1.15, 5.15, 0.85], shades=shades, keep_whole=False)
    total = sum(mn for _, _, mn, k in day if k not in ("lunch",))
    return total


def build():
    doc = Document()
    style_headings(doc)
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    for name in ("Heading 1", "Heading 2"):
        doc.styles[name].paragraph_format.keep_with_next = True
        doc.styles[name].paragraph_format.keep_together = True

    add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION,
                   conducted_by="Tertiary Infotech Pte Ltd",
                   org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                   course_logo=None, course_code=C.COURSE_CODE)

    h1(doc, "Document Version Control Record")
    add_version_control(doc, VERSIONS)

    h1(doc, "Table of Contents")
    para(doc, "Generated contents for the fixed v3.0 pagination.", size=9.5, color=GREY)
    table(doc, ["Section", "Page"],
          [["1. Course Overview", "4"],
           ["2. Daily Schedule", "6"],
           ["3. Activities", "8"],
           ["4. Tools and Resources", "8"],
           ["5. Assessment", "9"]],
          widths=[6.2, 0.95])
    doc.add_page_break()

    # ---- overview  (add_toc already ends the page; a second break renders blank)
    h1(doc, "1. Course Overview")
    para(doc, f"{C.TITLE} is a {C.DURATION} WSQ course that equips learners to identify, "
              "analyse and mitigate the security risks introduced by generative AI systems and "
              "by autonomous AI agents. The course uses documented cases and clearly labelled "
              "realistic simulations, grounded in current security and governance "
              "frameworks, including the OWASP Top 10 for LLM Applications (2026), the OWASP Top "
              "10 for Agentic Applications (2026), the NIST AI Risk Management Framework, MITRE "
              "ATLAS, the IMDA Model AI Governance Framework for Agentic AI, and Singapore's "
              "PDPA and current PDPC guidance. Proposed consultation material is not presented "
              "as final law or guidance.")
    para(doc, "The course covers AI security for BOTH generative AI (prompt injection, data "
              "leakage, poisoning, PDPA exposure) and autonomous agents (uncontrolled and "
              "destructive execution, tool misuse, agent vulnerabilities and cyber-attack "
              "chains).")

    h2(doc, "Skills Framework Alignment")
    table(doc, ["Element", "Detail"],
          [["TSC Title", C.TSC_TITLE],
           ["TSC Code", C.TSC_CODE],
           ["Proficiency Level", C.PROFICIENCY],
           ["Course Code", C.COURSE_CODE],
           ["Duration", C.DURATION],
           ["Delivery", "Instructor-led classroom with facilitated case-study activities"]],
          widths=[1.9, 5.25])
    para(doc, "")
    para(doc, "Note: the accredited knowledge and ability statements below are reproduced "
              "verbatim. They are taught and assessed as accredited; AI security is the "
              "delivery context through which each statement is evidenced.", size=9.5,
         color=GREY)

    h2(doc, "Learning Outcomes")
    table(doc, ["LO", "Learning Outcome"],
          [[c, d] for c, d in C.LEARNING_OUTCOMES], widths=[0.7, 6.45])

    h2(doc, "Knowledge Statements (assessed by the Written Assessment)")
    table(doc, ["Code", "Knowledge statement"],
          [[c, d] for c, d in C.TSC_KNOWLEDGE], widths=[0.7, 6.45])

    h2(doc, "Ability Statements (assessed by the Case Study)")
    table(doc, ["Code", "Ability statement"],
          [[c, d] for c, d in C.TSC_ABILITIES], widths=[0.7, 6.45])

    doc.add_page_break()

    # ---- schedule
    h1(doc, "2. Daily Schedule")
    para(doc, "Each training day runs 9:30am – 6:30pm and delivers 8 instructional hours. "
              "A 1-hour lunch break is excluded from instructional time; short tea breaks are "
              "counted within the day.", size=10, color=GREY)

    h2(doc, "Day 1 — Generative AI Security")
    para(doc, "Slides 1–98 of the trainer deck.", size=9.5, color=GREY)
    t1 = day_table(doc, DAY1)
    para(doc, "")
    para(doc, f"Day 1 instructional total: {t1} minutes ({t1/60:.0f} hours), excluding the "
              "1-hour lunch break.", size=10, bold=True)

    doc.add_page_break()
    h2(doc, "Day 2 — Autonomous Agent Security")
    para(doc, "Slides 99–207 of the trainer deck.", size=9.5, color=GREY)
    t2 = day_table(doc, DAY2)
    para(doc, "")
    para(doc, f"Day 2 instructional total: {t2} minutes ({t2/60:.0f} hours), excluding the "
              "1-hour lunch break.", size=10, bold=True)

    doc.add_page_break()

    # ---- activities
    h1(doc, "3. Activities")
    para(doc, "Every activity is either a clearly labelled realistic simulation or an "
              "evidence-based documented-case review. Each has its own folder under "
              "activities/ containing learner instructions, scenarios, discussion questions, "
              "checklists and printable PDFs. Full step-by-step facilitation detail is in the "
              "Learner Guide.")
    table(doc, ["#", "Activity", "Day", "Duration", "K/A assessed"],
          [["1", "Threat Modelling a Generative AI Concierge", "1", "45 min", "K2, K3, A4"],
           ["2", "Prompt Injection and the PDPA Breach Decision", "2", "60 min", "K4, K1"],
           ["3", "Selecting a Security Framework for GenAI and Agents", "2", "60 min", "A3, A5"],
           ["4", "Evidence-Based Rogue Agent Incident Review", "2", "60 min", "K5"],
           ["5", "Agent Governance and the Deployment Gate (capstone)", "2", "25 min", "A1, A2"]],
          widths=[0.4, 3.6, 0.5, 0.9, 1.75])

    h1(doc, "4. Tools and Resources")
    table(doc, ["Resource", "Purpose"],
          [["Trainer slide deck (PPTX/PDF)", "Facilitation, diagrams and case-study prompts"],
           ["Learner Guide", "Detailed step-by-step notes and activity walkthroughs"],
           ["Activity packs (activities/)", "Learner scenarios, discussion questions, checklists and evidence worksheets"],
           ["LMS/TMS portal", "lms-tms.tertiaryinfotech.com — course material and submission"],
           ["Whiteboard / flip chart", "Group threat modelling and kill-chain reconstruction"],
           ["OWASP LLM & ASI Top 10 (2026)", "Reference taxonomies used throughout"],
           ["IMDA / PDPC publications", "Singapore governance and data-protection reference"]],
          widths=[2.5, 4.65])

    doc.add_page_break()
    h1(doc, "5. Assessment")
    para(doc, "Assessment is conducted at the end of Day 2. The briefing for assessment is "
              "delivered before the assessment begins.")
    table(doc, ["Instrument", "Covers", "Detail"],
          [["Written Assessment (SAQ)", "K1 – K5", C.ASSESSMENT["wa"]],
           ["Case Study", "A1 – A5 via LO1 – LO3", C.ASSESSMENT["cs"]],
           ["Format", "—", C.ASSESSMENT["format"]],
           ["Grading", "—", C.ASSESSMENT["grading"]],
           ["Re-assessment", "—", "Available for learners assessed Not Yet Competent"]],
          widths=[1.75, 1.35, 4.05])
    para(doc, "")
    para(doc, "Funding eligibility requires a minimum 75% attendance recorded through SSG "
              "digital attendance, an assessment outcome of Competent, and completion of the "
              "TRAQOM survey.", size=10)

    add_page_numbers(doc)
    enable_update_fields(doc)

    assert t1 == 480, f"Day 1 is {t1} min, expected 480"
    assert t2 == 480, f"Day 2 is {t2} min, expected 480"

    out = os.path.join(HERE, f"Lesson Plan - {C.COURSE_CODE} - {C.TITLE}.docx")
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Day 1 = {t1} min · Day 2 = {t2} min")


if __name__ == "__main__":
    build()
