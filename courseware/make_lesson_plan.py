#!/usr/bin/env python3
"""WSQ Lesson Plan — AI Security for Autonomous AI Agents (TGS-2025060473).

1 day × 8 instructional hours (9:00am–6:00pm with a 1-hour lunch; tea breaks
counted within the day) plus a 1-hour assessment (6:00pm–7:00pm). Three topics
map onto LO1/LO2/LO3. Asserts 480 min of instruction before saving.
"""
import os, sys

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL = os.path.join(HERE, "..", ".claude", "skills", "tertiary-lesson-plan")
sys.path.insert(0, os.path.abspath(SKILL))
sys.path.insert(0, HERE)

import prodoc
# This course is delivered by Tertiary Infotech Pte Ltd (UEN 201200696W) — not the
# Academy entity hardcoded in the shared helper. Override before use.
prodoc.ORG = "Tertiary Infotech Pte Ltd"
prodoc.UEN = "UEN: 201200696W"
prodoc.COPYRIGHT = ("This material belongs to Tertiary Infotech Pte Ltd (UEN: 201200696W). "
                    "All Rights Reserved.")

from prodoc import (add_cover_page, add_version_control, add_toc, add_page_numbers,
                    enable_update_fields, style_headings, _shade_cell)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import course_data as C

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)
ASSETS = os.path.join(HERE, "assets")

VERSIONS = [
    ("1.0", "16 June 2025", "Initial release.", "Dr Alfred Ang"),
    ("2.0", "17 August 2026", "Retitled and rebuilt for autonomous-agent security; five activities "
                            "added. TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("2.1", "18 August 2026", "Added prompt-injection, PDPA, guardrail, approval, skill/plugin and "
                            "visual coverage. TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("3.0", "20 August 2026", "Evidence-grounded 207-slide rebuild with product boundaries, cases, "
                            "controls, activities and source register. TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.0", "25 August 2026", "Rebuilt as a 1-day course with three LO-aligned topics, eight no-code "
                            "activities and a reflection-based practical assessment. TSC K/A text "
                            "unchanged.", "Dr Alfred Ang"),
    ("4.1", "25 August 2026", "Refreshed trainer visuals and embedded the YouTube Online Video; "
                            "scope, timing and TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("4.2", "25 August 2026", "Added coding-harness and latest-AI-agents reference slides to the "
                            "deck (Claude Code, Codex, DeepSeek Harness; OpenClaw, Hermes, Prime "
                            "Agent, OpenWorker, QM). Schedule and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.3", "25 August 2026", "Beefed up Topic 3 with Singapore AI case studies (adoption vs "
                            "cybersecurity, gen-AI brand damage, agentic testing, healthcare "
                            "AgentSea), a brand-damage discussion/debrief, a 'Key Challenges of AI "
                            "to Business' capstone and a visual multi-agent map. Schedule and TSC "
                            "K/A text unchanged.", "Dr Alfred Ang"),
    ("4.4", "25 August 2026", "Activity 5 rebuilt around the AI Data Policy Generator website "
                            "(learner-supplied training API key, 7-part IMDA/PDPA-aligned policy "
                            "framework). Slides, LG and activity pack aligned; added a "
                            "worked-example slide with concrete Scope, Roles, Rules and Review "
                            "examples. Schedule timing and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.5", "25 August 2026", "Added Topic 1 slide 'Probabilistic Output — Why Traditional "
                            "Security Struggles' contrasting deterministic software with "
                            "probabilistic LLM output and the challenge to rule-based defences. "
                            "Schedule timing and TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("4.6", "25 August 2026", "Added Topic 1 slide 'When the Loop Goes Wrong — Misreading "
                            "the Goal' after the concrete agentic-loop example: an open-ended "
                            "'remove the malware at all costs' goal escalates to wiping every "
                            "file, motivating the Topic 3 guardrails. Schedule timing and TSC "
                            "K/A text unchanged.", "Dr Alfred Ang"),
    ("4.7", "25 August 2026", "Added Topic 1 slide 'An Agent on WhatsApp Is Exposed to the "
                            "Internet' after the TIA Support slide: a public messaging agent "
                            "accepts messages from anyone, so prompt injection and data "
                            "poisoning can extract confidential data or harm the server. "
                            "Schedule timing and TSC K/A text unchanged.", "Dr Alfred Ang"),
    ("4.8", "26 August 2026", "Added verified case study 'Case — A Malicious Skill on ClawHub' "
                            "(Snyk, Feb 2026) at the end of the Topic 1 Skills and Tools "
                            "section: a fake Google skill whose SKILL.md instructions tricked "
                            "users into installing malware. Schedule timing and TSC K/A text "
                            "unchanged.", "Dr Alfred Ang"),
    ("4.9", "26 August 2026", "Activity 3 reworked: the agent now creates a 10-slide 'AI Security "
                            "for AI Agents' deck (saved to the Downloads folder), then redesigns "
                            "it from an uploaded image template — replacing the Envato-template "
                            "animated-deck flow. Schedule timing and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.10", "26 August 2026", "Activity 2 reworked: learners now upload MOCK-MARKETING-DATA.xlsx "
                            "to the Hermes + MiniMax agent and prompt it to analyse the data, "
                            "generate charts and produce a .docx report with insights, analysis "
                            "and recommendations. Schedule timing and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.11", "26 August 2026", "Added four Topic 1 slides on applications of AI agents in "
                            "cyber security: threat detection (real-time detection, automated "
                            "vulnerability analysis, malware classification), network and "
                            "social-engineering security (intrusion detection, attack "
                            "classification, phishing defence, log anomaly detection), incident "
                            "response (patch generation, playbooks, root-cause analysis) and "
                            "Security Operations Centres (threat hunting, predictive analytics, "
                            "adaptive decision-making). Learner Guide notes renumbered. TSC K/A "
                            "text unchanged.", "Dr Alfred Ang"),
    ("4.12", "26 August 2026", "Added Topic 1 slide 'Agent-to-Agent Messaging - A Security Blind "
                            "Spot' after the multi-agent worked example: bad or malicious "
                            "output cascades agent-to-agent, and a growing topology "
                            "multiplies message paths - reducing visibility and making "
                            "data-flow tracing and attribution harder. Beefed up Topic 3 "
                            "security content with ten slides woven into the "
                            "existing sections: multi-layer bias mitigation; the four structural "
                            "security gaps of agents (multi-step inputs, tool chaining, opaque "
                            "execution, untrusted entities); direct vs indirect prompt injection; "
                            "jailbreaking and jailbreak persistence; a new 'Advanced Threats — A "
                            "Look Ahead' section (data poisoning, backdoor attacks, AI supply-chain "
                            "risk, agent-to-agent attacks, misalignment / Goodhart's Law); and "
                            "human-in-the-loop at machine speed. Schedule timing and TSC K/A text "
                            "unchanged.", "Dr Alfred Ang"),
    ("4.13", "26 August 2026", "Practical instrument renamed from Practical Performance to Case "
                            "Study across the assessment set, slides, Lesson Plan and Learner "
                            "Guide. Each Case Study task trimmed to TWO questions: Task 1 keeps "
                            "the generative-vs-agentic-vs-agent and useful-vs-not-trusted "
                            "questions; Task 2 keeps the weak-vs-strong prompt and "
                            "prompt-principles questions; Task 3 keeps the leaky-vs-guarded "
                            "chatbot and accountability questions. Marks rebalanced (20/24/26, "
                            "total 70). Schedule timing and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
    ("4.14", "26 August 2026", "Learning Outcomes slide now carries the full accredited LO1-LO3 "
                            "wording. Activity overview slides for Activities 1, 4, 6, 7 and 8 "
                            "regenerated with activity-specific step flows and 'You'll produce' "
                            "lines matching the activity packs and Learner Guide. Thank You "
                            "slide layout corrected. Activity PDFs re-stamped to the current "
                            "version. Schedule timing and TSC K/A text unchanged.",
     "Dr Alfred Ang"),
]

# (time, topic, minutes, kind)  kind: topic | activity | break | admin | assess
# One 8-instructional-hour day (480 min incl. tea breaks, excl. the 1-hour lunch)
# running 9:00am-6:00pm, followed by the 1-hour assessment 6:00pm-7:00pm.
DAY1 = [
    ("9:00 – 9:20", "Welcome, digital attendance, introductions, learning outcomes and ground rules", 20, "admin"),
    ("9:20 – 10:05", "Topic 1: Brief history of AI 2023-2026; how generative AI works — autoregressive LLM, training and inference (K2, K5)", 45, "topic"),
    ("10:05 – 10:45", "Topic 1: Real generative-AI use cases (AI fashion models, AI video, Klarna); context engineering (K3, A3, A4)", 40, "topic"),
    ("10:45 – 11:00", "Tea break", 15, "break"),
    ("11:00 – 11:45", "Topic 1: Agentic loop and harness engineering; AI agents, OpenClaw and Hermes, skills, tools and multi-agent systems (K2, K3)", 45, "topic"),
    ("11:45 – 12:30", "Activity 1: Talk to the TIA Support AI agent on WhatsApp and post reflections to the Pinboard; debrief (A2, A4)", 45, "activity"),
    ("12:30 – 1:30", "Lunch break", 60, "lunch"),
    ("1:30 – 1:45", "Topic 1 summary: generative AI vs agentic AI vs AI agents", 15, "topic"),
    ("1:45 – 2:20", "Topic 2: Install the Hermes agent on MiniMax M2.7; principles of prompt engineering, good vs bad prompts (K4, A5)", 35, "topic"),
    ("2:20 – 3:10", "Activity 2 & 3: Upload mock Excel data and prompt the agent into charts and a .docx insights report, then build a 10-slide PPT and redesign it from an uploaded image template (K4, A5)", 50, "activity"),
    ("3:10 – 3:25", "Tea break", 15, "break"),
    ("3:25 – 3:50", "Activity 4: Install tools and skills, re-run the Excel and PPT tasks; Topic 2 reflection and debrief (A1, A5)", 25, "activity"),
    ("3:50 – 4:35", "Topic 3: AI data governance and accountability (Moffatt v Air Canada); Activity 5 — generate an AI data governance policy with the policy-generator website (A1, A2)", 45, "activity"),
    ("4:35 – 5:15", "Topic 3: Job impact and redesign; Activity 6 — coaching role-play simulator (A2)", 40, "activity"),
    ("5:15 – 6:00", "Topic 3: AI-agent cybersecurity risks, advanced threats (prompt injection, jailbreaking, data poisoning, supply-chain and agent-to-agent attacks) and safe rollout; Activities 7 & 8 — leaky vs guarded chatbots and security reflection (A1, A2)", 45, "activity"),
    ("6:00 – 6:35", "Briefing for assessment; Written Assessment (SAQ) — 5 questions, one per K statement (K1-K5)", 35, "assess"),
    ("6:35 – 7:00", "Case Study — 3 reflection tasks (two questions each) on the activities completed; TRAQOM close", 25, "assess"),
]

SHADE = {"topic": "E8F0FE", "activity": "E6F7F1", "break": "FDF3E3", "lunch": "FDF3E3",
         "admin": "F2F4F7", "assess": "FDECEC"}


def h1(doc, t):
    p = doc.add_heading(t, level=1); return p


def h2(doc, t):
    return doc.add_heading(t, level=2)


def para(doc, t, size=11, bold=False, color=DARK, after=6):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    r.font.name = "Arial"
    p.paragraph_format.space_after = Pt(after)
    return p


def table(doc, headers, rows, widths=None, shades=None, keep_whole=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(htxt)
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Arial"
        _shade_cell(c, "1F6FEB")
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5); r.font.name = "Arial"
            r.font.bold = (j == 0 and shades is None)
        if shades:
            for c in cells:
                _shade_cell(c, shades[i])
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    for row_i, row in enumerate(t.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                if keep_whole:
                    p.paragraph_format.keep_with_next = row_i < len(t.rows) - 1
    return t


def day_table(doc, day):
    rows = [(tm, tp, f"{mn} min") for tm, tp, mn, _ in day]
    shades = [SHADE[k] for _, _, _, k in day]
    table(doc, ["Time", "Topic / Activity", "Duration"], rows,
          widths=[1.15, 5.15, 0.85], shades=shades, keep_whole=False)
    total = sum(mn for _, _, mn, k in day if k not in ("lunch", "assess"))
    return total


def build():
    doc = Document()
    style_headings(doc)
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    for name in ("Heading 1", "Heading 2"):
        doc.styles[name].paragraph_format.keep_with_next = True
        doc.styles[name].paragraph_format.keep_together = True

    add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION,
                   conducted_by="Tertiary Infotech Pte Ltd",
                   org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                   course_logo=None, course_code=C.COURSE_CODE)

    add_version_control(doc, VERSIONS)

    add_toc(doc)

    # ---- overview
    h1(doc, "1. Course Overview")
    para(doc, f"{C.TITLE} is a {C.DURATION} WSQ course that introduces generative AI, agentic "
              "AI and autonomous AI agents, and the security, governance and ethical risks they "
              "bring to customer-service and hospitality settings. Learners work entirely on "
              "ready-made websites, chatbots and AI agents — no coding is required. The day is "
              "organised as three topics mapped one-to-one onto the three learning outcomes.")
    para(doc, "Topic 1 (LO1) explains how generative AI works and how it becomes agentic — the "
              "agentic loop, harness engineering, AI agents such as OpenClaw and Hermes, skills, "
              "tools and multi-agent systems. Topic 2 (LO2) covers prompt engineering and "
              "post-training, using the Hermes agent on a MiniMax model to analyse data and "
              "build presentations. Topic 3 (LO3) covers AI data governance, job impact and "
              "redesign, and the cybersecurity risks of autonomous agents, including a safe "
              "roll-out framework. Real, dated examples and clearly labelled simulations are "
              "grounded in current governance references such as the IMDA Model AI Governance "
              "Framework, Singapore's PDPA and PDPC guidance, and the OWASP Top 10 for LLM and "
              "Agentic Applications.")

    h2(doc, "Skills Framework Alignment")
    table(doc, ["Element", "Detail"],
          [["TSC Title", C.TSC_TITLE],
           ["TSC Code", C.TSC_CODE],
           ["Proficiency Level", C.PROFICIENCY],
           ["Course Code", C.COURSE_CODE],
           ["Duration", C.DURATION],
           ["Delivery", "Instructor-led classroom with facilitated case-study activities"]],
          widths=[1.9, 5.25])
    para(doc, "")
    para(doc, "Note: the accredited knowledge and ability statements below are reproduced "
              "verbatim. They are taught and assessed as accredited; AI security is the "
              "delivery context through which each statement is evidenced.", size=9.5,
         color=GREY)

    h2(doc, "Learning Outcomes")
    table(doc, ["LO", "Learning Outcome"],
          [[c, d] for c, d in C.LEARNING_OUTCOMES], widths=[0.7, 6.45])

    h2(doc, "Knowledge Statements (assessed by the Written Assessment)")
    table(doc, ["Code", "Knowledge statement"],
          [[c, d] for c, d in C.TSC_KNOWLEDGE], widths=[0.7, 6.45])

    h2(doc, "Ability Statements (assessed by the Case Study)")
    table(doc, ["Code", "Ability statement"],
          [[c, d] for c, d in C.TSC_ABILITIES], widths=[0.7, 6.45])

    doc.add_page_break()

    # ---- schedule
    h1(doc, "2. Course Schedule")
    para(doc, "The training day runs 9:00am – 6:00pm and delivers 8 instructional hours. "
              "A 1-hour lunch break is excluded from instructional time; short tea breaks are "
              "counted within the day. The 1-hour assessment runs 6:00pm – 7:00pm.", size=10,
         color=GREY)

    h2(doc, "One-Day Schedule (Topics 1–3)")
    # Derive the slide count from the deck content so this line can never go stale.
    from v40_content import SLIDES as _DECK
    para(doc, f"Slides 1–{len(_DECK)} of the trainer deck.", size=9.5, color=GREY)
    t1 = day_table(doc, DAY1)
    para(doc, "")
    para(doc, f"Instructional total: {t1} minutes ({t1/60:.0f} hours), excluding the 1-hour "
              "lunch break and the 1-hour assessment.", size=10, bold=True)

    doc.add_page_break()

    # ---- activities
    h1(doc, "3. Activities")
    para(doc, "Every activity uses a ready-made website, chatbot or AI agent — learners do not "
              "write any code. Each has its own folder under activities/ containing learner "
              "instructions, prompt cards, mock data and printable PDFs. All data is fictional. "
              "Full step-by-step facilitation detail is in the Learner Guide.")
    table(doc, ["#", "Activity", "Topic", "Duration", "LO / A"],
          [["1", "Talk to an AI Agent (TIA Support WhatsApp) and reflect on the Pinboard", "1", "45 min", "LO1 · A2, A4"],
           ["2", "Analyse mock Excel marketing data into a charted .docx report with the Hermes + MiniMax agent", "2", "25 min", "LO2 · A5"],
           ["3", "Create a 10-slide PPT with the agent, then redesign it from an uploaded image template", "2", "25 min", "LO2 · A5"],
           ["4", "Install tools and skills, then re-run the Excel and PPT tasks", "2", "25 min", "LO2 · A1, A5"],
           ["5", "Generate an AI Data Governance Policy with the policy-generator website", "3", "30 min", "LO3 · A1, A2"],
           ["6", "Coach a worried team member (AI role-play simulator)", "3", "30 min", "LO3 · A2"],
           ["7", "Leaky vs guarded chatbot — see and stop a PII leak", "3", "30 min", "LO3 · A1, A2"],
           ["8", "Reflect on agent security and decide go / no-go", "3", "15 min", "LO3 · A1, A2"]],
          widths=[0.35, 3.65, 0.6, 0.85, 1.4])
    para(doc, "Times overlap with the topic teaching blocks in the schedule above; activity "
              "folders and the reflection Pinboard (alfredang.github.io/pinboard) are used "
              "throughout the day.", size=9.5, color=GREY)

    h1(doc, "4. Tools and Resources")
    table(doc, ["Resource", "Purpose"],
          [["Trainer slide deck (PPTX/PDF)", "Facilitation, diagrams and concept walkthroughs"],
           ["Learner Guide", "Detailed notes and step-by-step activity walkthroughs"],
           ["Activity packs (activities/)", "Ready-made websites/chatbots, prompt cards and mock data"],
           ["TIA Support on WhatsApp (+65 8866 6375)", "Live OpenClaw-powered AI agent for Activity 1"],
           ["Reflection Pinboard", "alfredang.github.io/pinboard — group reflections by risk theme"],
           ["Hermes Agent + MiniMax", "hermes-agent.nousresearch.com and minimax.io for Topic 2"],
           ["LMS/TMS portal", "lms-tms.tertiaryinfotech.com — course material and submission"],
           ["IMDA / PDPC / OWASP references", "Governance and security reference material"]],
          widths=[2.9, 4.25])

    doc.add_page_break()
    h1(doc, "5. Assessment")
    para(doc, "Assessment is conducted at the end of the day (6:00pm – 7:00pm). The briefing "
              "for assessment is delivered before the assessment begins.")
    table(doc, ["Instrument", "Covers", "Detail"],
          [["Written Assessment (SAQ)", "K1 – K5", C.ASSESSMENT["wa"]],
           ["Case Study", "LO1 – LO3 (A1 – A5)", C.ASSESSMENT["cs"]],
           ["Format", "—", C.ASSESSMENT["format"]],
           ["Grading", "—", C.ASSESSMENT["grading"]],
           ["Re-assessment", "—", "Available for learners assessed Not Yet Competent"]],
          widths=[1.75, 1.5, 3.9])
    para(doc, "")
    para(doc, "The Case Study asks learners to document their own observations and "
              "reflections from the activities they completed during the day, so it is grounded "
              "in what each learner actually did in class and aligned to the slides and labs.",
         size=10)
    para(doc, "Funding eligibility requires a minimum 75% attendance recorded through SSG "
              "digital attendance, an assessment outcome of Competent, and completion of the "
              "TRAQOM survey.", size=10)

    add_page_numbers(doc)
    enable_update_fields(doc)

    assert t1 == 480, f"Instructional total is {t1} min, expected 480"

    out = os.path.join(HERE, f"Lesson Plan - {C.COURSE_CODE} - {C.TITLE}.docx")
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Instruction = {t1} min")


if __name__ == "__main__":
    build()
